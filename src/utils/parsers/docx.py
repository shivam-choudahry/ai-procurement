"""
DOCX parser using python-docx.
Handles paragraphs and tables.
"""

import io
import logging

logger = logging.getLogger(__name__)


def parse_docx(content: bytes) -> str:
    """
    Extract text from a DOCX file, including tables.

    Args:
        content: Raw DOCX bytes

    Returns:
        Normalized plain text
    """
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        sections = []

        # Extract paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Preserve heading structure
                if para.style.name.startswith("Heading"):
                    sections.append(f"\n## {text}")
                else:
                    sections.append(text)

        # Extract tables
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                sections.append("\n[TABLE]\n" + "\n".join(rows) + "\n[/TABLE]")

        logger.info(f"python-docx: extracted {len(sections)} sections")
        return "\n".join(sections)

    except ImportError:
        logger.error("python-docx not installed")
        return "[DOCX parse error: python-docx not installed]"
    except Exception as e:
        logger.error(f"DOCX parse error: {e}")
        return f"[DOCX parse error: {e}]"

